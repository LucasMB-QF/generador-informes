from fastapi import FastAPI, UploadFile, File, Request, HTTPException
from fastapi.responses import Response, HTMLResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from openpyxl import load_workbook
from docx import Document
from docx.text.paragraph import Paragraph
import re
from io import BytesIO
from pathlib import Path
import logging
from urllib.parse import quote

# Configuración de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Rutas y carpetas
current_dir = Path(__file__).resolve().parent
templates_dir = current_dir / "templates"

app = FastAPI()

app.mount("/static", StaticFiles(directory="templates"), name="static")
templates = Jinja2Templates(directory=str(templates_dir))

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Regex para {{Hoja!Celda}} o {{campo}}
campo_regex = re.compile(r"\{\{\s*([^\{\}]+?)\s*\}\}")

# --- Formateo de valores ---

def formatear_valor(valor):
    if isinstance(valor, (float, int)):
        # Formatea con 1 decimal y reemplaza punto por coma
        return f"{valor:.1f}".replace(".", ",")
    return str(valor) if valor is not None else ""

# --- Lectura desde Excel ---

def obtener_valor(wb, hoja_nombre, celda):
    try:
        hoja = wb[hoja_nombre]
        valor = hoja[celda].value
        if valor is None:
            logger.warning(f"Celda vacía: {hoja_nombre}!{celda}")
        return formatear_valor(valor)
    except Exception as e:
        logger.error(f"Error en celda {hoja_nombre}!{celda}: {str(e)}")
        return ""

def obtener_valores_rango(wb, hoja_nombre, rango):
    try:
        hoja = wb[hoja_nombre]
        celdas = hoja[rango]
        fila = celdas[0]
        return [formatear_valor(c.value) for c in fila]
    except Exception as e:
        logger.error(f"Error en rango {hoja_nombre}!{rango}: {str(e)}")
        return []

# --- Reemplazo de campos en texto ---

def reemplazar_campos(texto, wb):
    def reemplazo(match):
        campo = match.group(1)
        if '!' in campo:
            hoja, celda_o_rango = campo.split('!', 1)
            hoja = hoja.strip()
            celda_o_rango = celda_o_rango.strip()
            if ':' in celda_o_rango:
                valores = obtener_valores_rango(wb, hoja, celda_o_rango)
                return ', '.join(valores)
            else:
                return obtener_valor(wb, hoja, celda_o_rango)
        return ""
    return campo_regex.sub(reemplazo, texto)

# --- Reemplazo en párrafos (versión robusta) ---

def reemplazar_en_parrafo(parrafo: Paragraph, wb):
    # Verificar si hay campos a reemplazar
    texto_total = "".join(run.text for run in parrafo.runs)
    if not campo_regex.search(texto_total):
        return
    
    # Obtener texto reemplazado
    texto_reemplazado = reemplazar_campos(texto_total, wb)
    
    # Conservar estilo de la primera run (excepto negrita)
    if parrafo.runs:
        primera_run = parrafo.runs[0]
        estilo_original = {
            'italic': primera_run.italic,
            'underline': primera_run.underline,
            'font': primera_run.font.name,
            'size': primera_run.font.size,
            'color': primera_run.font.color.rgb if primera_run.font.color else None
        }
        
        # Limpiar todas las runs
        for run in parrafo.runs:
            run.text = ""
        
        # Restaurar texto y formato
        primera_run.text = texto_reemplazado
        primera_run.bold = False  # Esto corrige específicamente el problema del negrita
        primera_run.italic = estilo_original['italic']
        primera_run.underline = estilo_original['underline']
        if estilo_original['font']:
            primera_run.font.name = estilo_original['font']
        if estilo_original['size']:
            primera_run.font.size = estilo_original['size']
        if estilo_original['color']:
            primera_run.font.color.rgb = estilo_original['color']

# --- Procesamiento de documento Word ---

def procesar_documento(doc, wb):
    for p in doc.paragraphs:
        reemplazar_en_parrafo(p, wb)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    reemplazar_en_parrafo(p, wb)

# --- Rutas FastAPI ---

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/procesar")
async def procesar(
    archivo_excel: UploadFile = File(...),
    archivo_word: UploadFile = File(...)
):
    try:
        logger.info("Iniciando procesamiento de archivos...")

        if not archivo_excel.filename.endswith(('.xlsx', '.xlsm')):
            raise HTTPException(400, "El archivo Excel debe ser .xlsx o .xlsm")
        if not archivo_word.filename.endswith('.docx'):
            raise HTTPException(400, "El archivo Word debe ser .docx")

        excel_content = await archivo_excel.read()
        word_content = await archivo_word.read()

        with BytesIO(excel_content) as excel_stream:
            wb = load_workbook(filename=excel_stream, data_only=True)

            with BytesIO(word_content) as word_stream:
                doc = Document(word_stream)
                procesar_documento(doc, wb)

                output_stream = BytesIO()
                doc.save(output_stream)
                output_stream.seek(0)

                logger.info("Procesamiento completado correctamente")

                nombre_base = archivo_word.filename.rsplit(".", 1)[0]
                nombre_generado = f"{nombre_base} (generado).docx"
                nombre_generado_seguro = quote(nombre_generado)

                return Response(
                    content=output_stream.getvalue(),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f'attachment; filename="{nombre_generado_seguro}"',
                        "Access-Control-Expose-Headers": "Content-Disposition"
                    }
                )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en el procesamiento: {str(e)}", exc_info=True)
        raise HTTPException(500, f"Error interno del servidor: {str(e)}")

# --- Página de error personalizada ---

@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    return templates.TemplateResponse(
        "error.html",
        {
            "request": request,
            "status_code": exc.status_code,
            "detail": exc.detail
        },
        status_code=exc.status_code
    )
